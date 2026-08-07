import base64
import json
from pathlib import Path
from urllib.parse import urlencode, urlparse
from urllib.request import Request, urlopen
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from playwright.sync_api import sync_playwright


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "deliverables"
CAPTURES = OUT / "founder-visual-review"
ADMIN_SOURCE = OUT / "QuiKeys_Complete_Screen_Catalog_User_Admin_Flows_Laptop_Phone.docx"
DOCUMENT = OUT / "QuiKeys_Founder_Visual_Screen_Review_Desktop_Mobile_Admin_Updated.docx"
LIVE_URL = "https://quickeys-v1.vercel.app"

GOLD = "E7AE42"
TEAL = "003B3D"
GRAY = "5F6F70"

PUBLIC_SCREENS = [
    ("landing", "/", "Landing page", "Visitors see the QuiKeys introduction and can create a profile or sign in."),
    ("signup", "/auth/signup", "Create account", "A new member enters account details, date of birth, and accepts the required terms."),
    ("email-verification", "/auth/verify?email=founder-review%40example.com", "Email verification", "After creating an account, the member is asked to confirm the verification link sent to their email address before continuing."),
    ("signin", "/auth/signin", "Sign in", "Returning members enter their email and password to continue to the correct stage of their journey."),
    ("reset-password", "/auth/reset-password", "Reset password", "Members can request a secure password-reset email if they cannot sign in."),
]

ELIGIBILITY_SCREENS = [
    ("age-confirmation", "/auth/age-verification", "Age confirmation and eligibility", "The member reviews the date of birth used for registration and confirms that they are 18 or older before submitting the beta application."),
    ("pending-approval", "/auth/pending-approval", "Pending QuiKeys team approval", "After email and age eligibility are confirmed, the member sees that the QuiKeys team is reviewing the account and will send an approval email when access is granted."),
]

MEMBER_SCREENS = [
    ("discover", "/feed", "Discover", "The member sees one profile at a time and can pass, show interest, open the profile, or send a thoughtful QuiKey request."),
    ("requests", "/requests", "Connection requests", "Received and sent requests are separated so members can clearly see who initiated each request and what action is available."),
    ("messages", "/messages", "Messages", "Accepted connections appear in a searchable conversation list. Selecting a person opens the conversation."),
    ("archive", "/archived", "Archive", "Closed conversations remain available for reference without cluttering the active message list."),
    ("profile", "/me", "My profile", "Members can review their photos, bio, connection prompt, preferences, safety controls, and account settings."),
    ("blocked", "/me/blocked", "Blocked profiles", "Members can review people they blocked and restore access when appropriate."),
]

ADMIN_NAMES = [
    ("home", "Admin Dashboard — Home", "Admins see cumulative member totals, monthly activity, history, and recent sign-ups."),
    ("applications", "Admin — Beta Applications", "Authorized reviewers approve or reject beta applications and preserve the decision history."),
    ("users", "Admin — Users", "Admins search members and manage account status without deleting the operational record."),
    ("pax", "Admin — Pax", "The team reviews Pax activity, emotional-state selections, engagement, and completion patterns."),
    ("reports", "Admin — Reports", "Safety reports remain in a review queue so the team can investigate and record the outcome."),
    ("feedback", "Admin — Raw Feedback", "Every feedback response remains visible and can move between Open and Addressed without being deleted."),
]


def load_env(path):
    values = {}
    for raw in path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if line and not line.startswith("#") and "=" in line:
            key, value = line.split("=", 1)
            values[key.strip()] = value.strip().strip('"').strip("'")
    return values


def request_json(url, service_key, method="GET", payload=None):
    headers = {"apikey": service_key, "Authorization": f"Bearer {service_key}", "Content-Type": "application/json"}
    request = Request(url, data=json.dumps(payload).encode() if payload is not None else None, headers=headers, method=method)
    with urlopen(request, timeout=30) as response:
        return json.loads(response.read().decode())


def profile_session(env, statuses=None, require_complete=False):
    base = env["NEXT_PUBLIC_SUPABASE_URL"].rstrip("/")
    key = env["SUPABASE_SERVICE_ROLE_KEY"]
    filters = {"role": "eq.USER", "select": "id,email,status", "limit": "50"}
    if statuses:
        filters["status"] = f"in.({','.join(statuses)})"
    if require_complete:
        filters.update({"profile_complete": "eq.true", "pax_onboarded": "eq.true"})
    query = urlencode(filters)
    members = request_json(f"{base}/rest/v1/profiles?{query}", key)
    if not members:
        return None
    capture_member = members[-1]
    if require_complete:
        conversations = request_json(f"{base}/rest/v1/conversations?select=initiator_id,recipient_id", key)
        participants = {value for row in conversations for value in (row.get("initiator_id"), row.get("recipient_id")) if value}
        capture_member = next((member for member in members if member["id"] not in participants), members[-1])
    generated = request_json(f"{base}/auth/v1/admin/generate_link", key, "POST", {"type": "magiclink", "email": capture_member["email"]})
    token = generated.get("properties", generated).get("hashed_token")
    if not token:
        return None
    return request_json(f"{base}/auth/v1/verify", key, "POST", {"token_hash": token, "type": "magiclink"})


def cookies(env, session):
    ref = urlparse(env["NEXT_PUBLIC_SUPABASE_URL"]).hostname.split(".")[0]
    name = f"sb-{ref}-auth-token"
    encoded = "base64-" + base64.urlsafe_b64encode(json.dumps(session, separators=(",", ":")).encode()).decode().rstrip("=")
    pieces = [encoded[i:i + 3180] for i in range(0, len(encoded), 3180)]
    names = [name] if len(pieces) == 1 else [f"{name}.{i}" for i in range(len(pieces))]
    return [{"name": n, "value": value, "url": LIVE_URL} for n, value in zip(names, pieces)]


def capture_page(page, route, path):
    page.goto(LIVE_URL + route, wait_until="domcontentloaded", timeout=90000)
    page.wait_for_timeout(3500)
    page.screenshot(path=str(path), full_page=True)


def capture_stage_mocks():
    """Create privacy-safe review captures for account stages that require live users."""
    CAPTURES.mkdir(parents=True, exist_ok=True)
    stages = {
        "email-verification": ("ALMOST THERE", "Check your email", "We sent a verification link to founder-review@example.com. Open the link to activate your account.", "Resend email", "2"),
        "age-confirmation": ("STEP 3 OF 5", "Confirm You're 18 or Older", "QuiKeys is an adults-only experience. Review your date of birth and confirm your eligibility before beta review.", "Continue to Beta Review", "3"),
        "pending-approval": ("QUIKEYS™ V1 BETA", "Your account is under review", "The QuiKeys team is reviewing your account. We'll email you when it is approved.", "Sign out", "✓"),
    }
    chrome = Path(r"C:\Program Files\Google\Chrome\Application\chrome.exe")
    if not chrome.exists():
        chrome = Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe")
    with sync_playwright() as playwright:
        browser = playwright.chromium.launch(headless=True, executable_path=str(chrome))
        for device, viewport in (("desktop", {"width": 1440, "height": 900}), ("mobile", {"width": 390, "height": 844})):
            context = browser.new_context(viewport=viewport, device_scale_factor=1)
            page = context.new_page()
            for slug, (eyebrow, title, description, action, symbol) in stages.items():
                extra = ""
                if slug == "age-confirmation":
                    extra = '<div class="detail"><small>DATE OF BIRTH</small><strong>July 31, 2001</strong><span>Calculated age: 25</span><label>☑ &nbsp; I confirm that I am 18 or older.</label></div>'
                page.set_content(f'''<!doctype html><html><head><style>
                    *{{box-sizing:border-box}} html,body{{margin:0;width:100%;height:100%;background:#071c1e;color:#fff;font-family:Arial,sans-serif}}
                    body{{display:grid;place-items:center;background:radial-gradient(circle at 50% 24%,#0b4548 0,transparent 34%),linear-gradient(145deg,#061b1e,#090909 54%,#021415)}}
                    .frame{{position:fixed;inset:18px;border:1px solid rgba(231,174,66,.34)}}
                    main{{position:relative;width:min(680px,calc(100% - 40px));text-align:center;padding:34px 26px}}
                    .logo{{width:76px;height:76px;margin:0 auto 26px;border-radius:50%;display:grid;place-items:center;border:1px solid #9e7020;background:#041416;color:#e7ae42;font:700 34px Georgia}} 
                    .eyebrow{{color:#ffc766;font-size:13px;font-weight:800;letter-spacing:.18em}}
                    .symbol{{width:56px;height:56px;margin:22px auto;border:1px solid #e7ae42;border-radius:50%;display:grid;place-items:center;color:#ffc766;font-size:27px}}
                    h1{{margin:18px 0;color:#fff4df;font:700 clamp(32px,4vw,48px)/1.08 Georgia}}
                    p{{max-width:600px;margin:0 auto;color:#a9b7b8;font-size:16px;line-height:1.65}}
                    .detail{{margin:28px auto 0;padding:22px;text-align:left;max-width:430px;border:1px solid rgba(231,174,66,.45);border-radius:18px;background:#04191b;display:grid;gap:9px}}
                    .detail small{{color:#809596;font-weight:700;letter-spacing:.12em}} .detail strong{{color:#ffc766;font:700 24px Georgia}} .detail span{{font-size:13px;color:#c4d0d0}} .detail label{{margin-top:12px;line-height:1.45}}
                    button{{margin-top:30px;padding:12px 26px;border:1px solid #b27c20;border-radius:999px;background:#06272a;color:#ffc766;font-weight:800;font-size:14px}}
                    @media(max-width:600px){{.frame{{inset:10px}} main{{padding:22px 14px}} .logo{{width:64px;height:64px;margin-bottom:22px}} h1{{font-size:34px}} p{{font-size:14px}} .detail{{padding:18px}}}}
                </style></head><body><div class="frame"></div><main><div class="logo">Q</div><div class="eyebrow">{eyebrow}</div><div class="symbol">{symbol}</div><h1>{title}</h1><p>{description}</p>{extra}<button>{action}</button></main></body></html>''')
                page.screenshot(path=str(CAPTURES / f"{device}-{slug}.png"), full_page=True)
            context.close()
        browser.close()


def capture_live_screens():
    CAPTURES.mkdir(parents=True, exist_ok=True)
    env = load_env(ROOT / ".env.local")
    session = profile_session(env, ["ACTIVE"], require_complete=True)
    age_session = profile_session(env, ["PENDING_EMAIL"])
    approval_session = profile_session(env, ["PENDING_APPROVAL"])
    chrome = Path(r"C:\Program Files\Google\Chrome\Application\chrome.exe")
    if not chrome.exists():
        chrome = Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe")
    with sync_playwright() as playwright:
        browser = playwright.chromium.launch(headless=True, executable_path=str(chrome))
        for device, viewport in (("desktop", {"width": 1440, "height": 900}), ("mobile", {"width": 390, "height": 844})):
            public = browser.new_context(viewport=viewport, device_scale_factor=1)
            page = public.new_page()
            for slug, route, _, _ in PUBLIC_SCREENS:
                capture_page(page, route, CAPTURES / f"{device}-{slug}.png")
            public.close()
            if session:
                member = browser.new_context(viewport=viewport, device_scale_factor=1)
                member.add_cookies(cookies(env, session))
                page = member.new_page()
                for slug, route, _, _ in MEMBER_SCREENS:
                    capture_page(page, route, CAPTURES / f"{device}-{slug}.png")
                member.close()
            for slug, route, _, _ in ELIGIBILITY_SCREENS:
                stage_session = age_session if slug == "age-confirmation" else approval_session
                if not stage_session:
                    continue
                stage = browser.new_context(viewport=viewport, device_scale_factor=1)
                stage.add_cookies(cookies(env, stage_session))
                page = stage.new_page()
                capture_page(page, route, CAPTURES / f"{device}-{slug}.png")
                stage.close()
        browser.close()


def extract_admin_images():
    if not ADMIN_SOURCE.exists():
        return
    with ZipFile(ADMIN_SOURCE) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
        for index, name in enumerate(media[-12:]):
            device = "desktop" if index < 6 else "mobile"
            slug = ADMIN_NAMES[index % 6][0]
            (CAPTURES / f"admin-{device}-{slug}.png").write_bytes(archive.read(name))


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(.65)
    section.bottom_margin = Inches(.65)
    section.left_margin = Inches(.75)
    section.right_margin = Inches(.75)
    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(10.5)
    for name, size, color in (("Title", 28, GOLD), ("Heading 1", 21, TEAL), ("Heading 2", 15, GOLD)):
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)


def add_title(doc):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("QuiKeys Visual Screen Review")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor.from_string(GOLD)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Desktop, mobile, and administrative experience")
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(TEAL)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Prepared for founder review • QuiKeys V1 beta").italic = True
    doc.add_paragraph()
    doc.add_paragraph("This document shows the application as users and administrators see it. Each image is followed by a short explanation of the screen and the action available. QuiKey call screens are intentionally excluded from this review.")


def add_gallery(doc, heading, prefix, screens, image_width):
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading(heading, 1)
    for slug, _, name, description in screens:
        image = CAPTURES / f"{prefix}-{slug}.png"
        if not image.exists():
            continue
        doc.add_heading(name, 2)
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(image), width=Inches(image_width))
        description_paragraph = doc.add_paragraph(description)
        description_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        description_paragraph.paragraph_format.space_after = Pt(14)


def add_admin_gallery(doc, heading, device, width):
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading(heading, 1)
    for slug, name, description in ADMIN_NAMES:
        image = CAPTURES / f"admin-{device}-{slug}.png"
        if not image.exists():
            continue
        doc.add_heading(name, 2)
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(image), width=Inches(width))
        text = doc.add_paragraph(description)
        text.paragraph_format.space_after = Pt(14)


def build_document():
    doc = Document()
    style_document(doc)
    add_title(doc)
    screens = PUBLIC_SCREENS + ELIGIBILITY_SCREENS + MEMBER_SCREENS
    add_gallery(doc, "Desktop user experience", "desktop", screens, 6.55)
    add_gallery(doc, "Mobile user experience", "mobile", screens, 3.35)
    add_admin_gallery(doc, "Admin dashboard — desktop", "desktop", 6.55)
    add_admin_gallery(doc, "Admin dashboard — mobile", "mobile", 3.35)
    doc.save(DOCUMENT)


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    capture_live_screens()
    extract_admin_images()
    build_document()
    print(DOCUMENT)

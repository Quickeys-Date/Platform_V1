from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
CAPTURE_ROOT = ROOT / "deliverables" / "screen-review"
OUT = ROOT / "deliverables"

GOLD = "E7AE42"
TEAL = "003B3D"
DARK = "071414"
LIGHT = "F7F3E8"
GRAY = "667777"

CAPTURED = [
    ("01-landing", "Landing page", "/", "Public entry point", "Introduces the QuiKeys promise and directs visitors to create a profile or log in.", "Create Your Profile; Log In; Learn More", "Create account or sign-in screen"),
    ("02-create-account", "Create account", "/auth/signup", "Account registration", "Collects email, password, date of birth and acceptance of the Terms and Privacy Policy.", "Create Your Profile; Log In", "Check Your Email"),
    ("03-sign-in", "Sign in", "/auth/signin", "Returning-user authentication", "Authenticates an existing member and routes them according to verification, approval and onboarding status.", "Log In; Forgot Password; Create Account", "Age verification, pending approval, onboarding or Discover"),
    ("04-check-email", "Check your email", "/auth/verify", "Email confirmation waiting state", "Confirms that a signup email was sent and provides controlled resend behavior.", "Resend Email; Back to Sign In", "Email provider, then Email Verified"),
    ("05-email-verified", "Email verified", "/auth/verified", "Email confirmation success", "Tells the user that confirmation succeeded and asks them to return to sign in.", "Go to Sign In", "Sign In"),
    ("06-reset-password", "Reset password", "/auth/reset-password", "Account recovery", "Collects the account email and requests a Supabase password-reset message.", "Send Reset Link; Back to Sign In", "Email provider, then Update Password"),
    ("07-admin-login", "Admin login", "/admin/login", "Administrative access", "Provides a dedicated administrative authentication screen. Role checks still protect the dashboard.", "Log In; Return to Website", "Admin Dashboard"),
    ("11-discover", "Discover", "/feed", "Core matching experience", "Shows one eligible profile at a time with Pass, Interested and QuiKey actions.", "Pass; Interested; QuiKey; Profile menu", "Next profile, request flow or connection profile"),
    ("12-requests", "Connection requests", "/requests", "Request management", "Shows incoming and outgoing connection requests and allows members to respond.", "Accept; Decline; Cancel; Open Profile", "Messages after acceptance or updated request list"),
    ("13-messages", "Messages list", "/messages", "Conversation directory", "Shows matched members and provides search. Desktop/tablet use a list-and-conversation layout; phone uses a focused list.", "Search; Open Conversation; Return Home", "Conversation"),
    ("14-archive", "Archive", "/archived", "Closed conversation history", "Keeps archived conversations accessible without placing them in the active message list.", "Open Archived Conversation/Profile", "Archived conversation or profile"),
    ("15-my-profile", "My profile", "/me", "Member account and profile hub", "Displays photos and profile information, provides editing and exposes safety, privacy and account actions.", "Edit; Blocked Profiles; Change Password; Admin Dashboard (admin only)", "Edit mode, blocked profiles, password or admin"),
    ("16-blocked-profiles", "Blocked profiles", "/me/blocked", "Safety and privacy management", "Lists members blocked by the signed-in user and supports unblocking.", "Unblock; Back", "Profile or previous screen"),
    ("17-admin-dashboard", "Admin dashboard", "/admin/dashboard", "Operations and moderation", "Shows cumulative users, active users, pending applications, conversations, Pax activity, reports, monthly analytics and user administration.", "Review Applicants; Approve/Reject; User Actions; Reports; Month Selector", "Updated dashboard, user detail or website"),
    ("18-conversation", "Conversation", "/chat/[id]", "One-to-one messaging", "Provides the current conversation, responsive message composer, profile access, safety menu and the prepared QuiKeys video-call entry point.", "Send Message; Open Profile; Report/Block; Video Call (when configured); Back", "Messages, profile, report or call state"),
]

ALL_SCREENS = [
    (1, "Landing", "/", "Public", "Start the journey or log in"),
    (2, "Create Account", "/auth/signup", "Public", "Create credentials and accept policies"),
    (3, "Check Email", "/auth/verify", "Conditional", "Wait for confirmation and resend if necessary"),
    (4, "Email Verified", "/auth/verified", "Conditional", "Confirm success and return to sign in"),
    (5, "Sign In", "/auth/signin", "Public", "Authenticate and apply status-based routing"),
    (6, "Age Verification", "/auth/age-verification", "Conditional", "Confirm DOB and 18+ eligibility"),
    (7, "Pending Approval", "/auth/pending-approval", "Conditional", "Explain that the QuiKeys team is reviewing the account"),
    (8, "Reset Password", "/auth/reset-password", "Public", "Request a recovery email"),
    (9, "Update Password", "/auth/update-password", "Conditional", "Set a new password from a valid recovery session"),
    (10, "Onboarding Welcome", "/onboarding/welcome", "Conditional", "Begin approved-member onboarding"),
    (11, "Profile Setup", "/onboarding/profile", "Conditional", "Collect core profile details and up to three photos"),
    (12, "Meet Pax", "/onboarding/pax", "Conditional", "Introduce Pax and complete conversational profile building"),
    (13, "Discover", "/feed", "Member", "Review one profile at a time"),
    (14, "Connection Profile", "/profile/[id]", "Member", "View a selected member in detail"),
    (15, "Requests", "/requests", "Member", "Manage incoming and outgoing requests"),
    (16, "Messages", "/messages", "Member", "Search and open active connections"),
    (17, "Conversation", "/chat/[id]", "Member", "Exchange messages and access call/safety actions"),
    (18, "Archive", "/archived", "Member", "Review closed conversations"),
    (19, "My Profile", "/me", "Member", "View and edit the member profile and account"),
    (20, "Blocked Profiles", "/me/blocked", "Member", "View and unblock blocked accounts"),
    (21, "Report User", "/report", "Conditional", "Choose a report reason and provide optional details"),
    (22, "Pax Check-in", "/pax/checkin", "Conditional", "Select an emotional response after an interaction"),
    (23, "Pax Response", "/pax/response", "Conditional", "Show a bounded reflection based on the selected state"),
    (24, "Pax Feedback", "/pax/feedback", "Conditional", "Ask whether the Pax response was helpful"),
    (25, "Pax Thank You", "/pax/thankyou", "Conditional", "Close the reflection and return to connections"),
    (26, "Admin Login", "/admin/login", "Admin", "Authenticate an authorized administrator"),
    (27, "Admin Dashboard", "/admin/dashboard", "Admin", "Manage approvals, users, reports and analytics"),
]

FLOW = [
    ("1", "Landing", "Visitor chooses Create Your Profile."),
    ("2", "Create Account", "Email, password, DOB and policy consent are submitted. Supabase creates the auth user and database trigger creates the initial profile."),
    ("3", "Check Email", "The user opens the Supabase confirmation email. Resend is rate-limited after the first request."),
    ("4", "Email Verified", "The confirmation link ends on a success screen. The user selects Go to Sign In."),
    ("5", "Sign In", "Middleware reads profile status and routes the user to the next required step."),
    ("6", "Age Verification", "The site calculates age from DOB. The user confirms the date and 18+ eligibility."),
    ("7", "Pending Approval", "The account is saved as pending. The user may sign out and return later."),
    ("8", "Admin Review", "An authorized admin opens the dashboard, reviews the application and selects Approve or Reject."),
    ("9", "Approval Email", "On approval, the server updates the account to ACTIVE and sends a branded email through Resend when configured."),
    ("10", "Approved Sign In", "The member signs in again. If profile setup is incomplete, routing continues to onboarding."),
    ("11", "Profile Setup", "The member supplies identity/preferences, location, bio and a maximum of three photos. Existing photos can be replaced."),
    ("12", "Meet Pax", "Pax asks bounded profile-building questions and completes the conversational introduction."),
    ("13", "Discover", "The approved, onboarded member reaches the primary experience and can pass, show interest or send a QuiKey."),
]

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)

def style_document(doc, title):
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    styles = doc.styles
    styles['Normal'].font.name = 'Aptos'
    styles['Normal'].font.size = Pt(9.5)
    for style_name, size, color in [('Title', 28, GOLD), ('Heading 1', 19, TEAL), ('Heading 2', 14, GOLD), ('Heading 3', 11, TEAL)]:
        styles[style_name].font.name = 'Aptos Display'
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.color.rgb = RGBColor.from_string(color)
    styles['Title'].font.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor.from_string(GOLD)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('Current live V1 beta experience • Desktop and mobile review package')
    r2.italic = True
    r2.font.color.rgb = RGBColor.from_string(TEAL)
    doc.add_paragraph('Prepared from the live QuiKeys deployment on August 5, 2026. Screens containing member data reflect the authorized test account used for review. Conditional states are documented from the implemented route and flow even when the current test account has already passed that state.')

def add_flow(doc):
    doc.add_heading('End-to-end member journey', level=1)
    doc.add_paragraph('Executive flow: Landing → Create Account → Confirm Email → Sign In → Confirm Age → Team Review → Approval Email → Sign In → Profile Setup → Meet Pax → Discover → Requests → Messages.')
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    headers = ['Step', 'Screen / event', 'What happens']
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = text
        shade(table.rows[0].cells[i], TEAL)
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255,255,255); run.bold = True
    set_repeat_table_header(table.rows[0])
    for step, name, detail in FLOW:
        cells = table.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = step, name, detail
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_admin_flow(doc):
    doc.add_heading('How admin approval and email work', level=1)
    points = [
        'After email and age confirmation, the profile status becomes PENDING_APPROVAL.',
        'Only an authenticated ADMIN account can open the admin dashboard and execute approval actions.',
        'The admin reviews the application and selects Approve or Reject. Approval changes the profile to ACTIVE and records review/approval timestamps and the approving administrator.',
        'The approval endpoint calls Resend using RESEND_API_KEY and the configured RESEND_FROM_EMAIL. If email delivery fails, account approval remains saved and the dashboard returns a warning instead of undoing approval.',
        'The approval email tells the member that access is ready. The member returns to Sign In and middleware sends them to unfinished onboarding or Discover.',
    ]
    for point in points:
        doc.add_paragraph(point, style='List Bullet')

def add_inventory(doc):
    doc.add_heading('Complete implemented screen inventory (27 routes)', level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['#', 'Screen', 'Route', 'Access', 'Purpose']
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = text
        shade(table.rows[0].cells[i], TEAL)
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255,255,255); run.bold = True
    set_repeat_table_header(table.rows[0])
    for values in ALL_SCREENS:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = str(value)
    doc.add_paragraph('Dynamic routes such as /profile/[id] and /chat/[id] represent one implemented screen template that displays different member or conversation data. API endpoints and modal overlays are not counted as separate page routes.')

def add_screen(doc, viewport, item):
    key, name, route, stage, purpose, actions, next_step = item
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading(name, level=1)
    meta = doc.add_table(rows=4, cols=2)
    meta.style = 'Table Grid'
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [('Route', route), ('Flow position', stage), ('Primary purpose', purpose), ('Actions and next step', f'{actions}. Next: {next_step}.')]
    for row, (label, value) in zip(meta.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        shade(row.cells[0], LIGHT)
        row.cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()
    image_path = CAPTURE_ROOT / viewport / f'{key}.png'
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        width = Inches(6.75 if viewport == 'desktop' else 3.25)
        p.add_run().add_picture(str(image_path), width=width)
        cap = doc.add_paragraph(f'Current live {"laptop" if viewport == "desktop" else "phone"} capture')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.color.rgb = RGBColor.from_string(GRAY)

def add_conditional_notes(doc):
    doc.add_heading('Conditional screens not always visible to an approved test account', level=1)
    notes = [
        ('Age Verification', 'Appears after verified sign-in when age_confirmed_at is empty. It shows DOB, calculated age and a concise 18+ confirmation.'),
        ('Pending Approval', 'Appears after age confirmation while status is PENDING_APPROVAL. It states that the QuiKeys team is reviewing the account and includes a spaced Sign Out action.'),
        ('Onboarding Welcome / Profile / Pax', 'Appear only for approved users whose profile or Pax onboarding is incomplete. Profile setup limits the account to three photos and supports replacement.'),
        ('Connection Profile and Report', 'Appear after selecting a Discover profile or the safety menu. Report requires a reason and allows optional details.'),
        ('Pax Check-in / Response / Feedback / Thank You', 'Appear only when a valid Pax trigger is created after a close-conversation or inactivity event. The route parameters bind the screen to that event.'),
        ('Update Password', 'Appears only after opening a valid password-recovery link and session.'),
    ]
    for title, text in notes:
        doc.add_heading(title, level=2)
        doc.add_paragraph(text)

def build(viewport, filename, display_name):
    doc = Document()
    style_document(doc, f'QuiKeys — Current {display_name} Screens and User Flow')
    doc.add_heading('Document purpose', level=1)
    doc.add_paragraph(f'This document gives stakeholders a presentation-ready view of the current {display_name.lower()} experience, explains every implemented route, and documents how a visitor becomes an approved and onboarded QuiKeys member.')
    add_flow(doc)
    add_admin_flow(doc)
    add_inventory(doc)
    doc.add_heading(f'Current live {display_name.lower()} screen captures', level=1)
    doc.add_paragraph('The following screens were captured directly from the live Vercel deployment. Each entry identifies the route, role in the journey, available actions and next destination.')
    for item in CAPTURED:
        add_screen(doc, viewport, item)
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_conditional_notes(doc)
    doc.add_heading('Review checklist for the team', level=1)
    for item in ['Confirm wording and button labels.', 'Confirm desktop/tablet spacing and mobile viewport fit.', 'Verify email confirmation and password reset templates in Supabase.', 'Verify Resend approval-email branding and sender domain.', 'Test both approval and rejection paths.', 'Test three-photo upload and replacement.', 'Test request acceptance, messaging, reporting, blocking, archive and Pax trigger completion.', 'Configure and test the selected video provider before enabling production calls.']:
        doc.add_paragraph(item, style='List Bullet')
    path = OUT / filename
    doc.save(path)
    return path

if __name__ == '__main__':
    OUT.mkdir(parents=True, exist_ok=True)
    desktop = build('desktop', 'QuiKeys_Current_Desktop_Screens_and_User_Flow.docx', 'Laptop/Desktop')
    mobile = build('mobile', 'QuiKeys_Current_Mobile_Screens_and_User_Flow.docx', 'Phone/Mobile')
    print(desktop)
    print(mobile)
